"""
Resume Templates Module - DOCX Format
Pre-designed professional resume templates for users to download and customize
Generates editable DOCX files instead of PDFs
"""

from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class ResumeTemplateGenerator:
    """
    Generate professional resume templates in DOCX format for users to download and edit
    """
    
    TEMPLATES = {
        'professional_chronological': {
            'name': 'Professional Chronological',
            'description': 'Traditional format highlighting work experience. Best for professionals with 3+ years of experience.',
            'best_for': 'Experienced Professionals',
            'ats_friendly': True,
            'sections': ['Contact', 'Professional Summary', 'Experience', 'Skills', 'Education', 'Certifications']
        },
        'modern_skills_focused': {
            'name': 'Modern Skills-Focused',
            'description': 'Emphasizes technical skills and projects. Ideal for developers and tech professionals.',
            'best_for': 'Tech & Engineering Roles',
            'ats_friendly': True,
            'sections': ['Contact', 'Skills', 'Technical Projects', 'Experience', 'Education']
        },
        'entry_level_student': {
            'name': 'Entry-Level Student',
            'description': 'Perfect for students, recent graduates, and career changers with limited work experience.',
            'best_for': 'Students & New Graduates',
            'ats_friendly': True,
            'sections': ['Contact', 'Education', 'Skills', 'Projects', 'Internships/Experience', 'Activities']
        }
    }
    
    def generate_template(self, template_type: str, output_path: str) -> bool:
        """
        Generate a resume template DOCX file
        
        Args:
            template_type: Type of template ('professional_chronological', 'modern_skills_focused', 'entry_level_student')
            output_path: Path to save the DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            if template_type not in self.TEMPLATES:
                logger.error(f"Unknown template type: {template_type}")
                return False
            
            # Create document
            doc = Document()
            
            # Set narrow margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Generate template-specific content
            if template_type == 'professional_chronological':
                self._generate_professional_chronological(doc)
            elif template_type == 'modern_skills_focused':
                self._generate_modern_skills_focused(doc)
            elif template_type == 'entry_level_student':
                self._generate_entry_level_student(doc)
            
            # Save document
            doc.save(output_path)
            
            logger.info(f"✓ Template generated: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to generate template: {e}")
            return False
    
    def _add_name(self, doc, name="YOUR FULL NAME"):
        """Add centered name heading"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 41, 55)  # Dark gray
    
    def _add_contact(self, doc, contact_text):
        """Add centered contact information"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)  # Gray
    
    def _add_section_header(self, doc, title):
        """Add section header with border"""
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 41, 55)  # Dark gray
        # Add bottom border effect with a line of underscores
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(12)
    
    def _add_placeholder(self, doc, text, italic=True):
        """Add placeholder text"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        if italic:
            run.italic = True
        run.font.color.rgb = RGBColor(107, 114, 128)  # Gray
    
    def _add_bullet(self, doc, text):
        """Add bullet point"""
        p = doc.add_paragraph(text, style='List Bullet')
        run = p.runs[0]
        run.font.size = Pt(10)
    
    def _add_job_title(self, doc, title, company, location):
        """Add job title with company and location"""
        p = doc.add_paragraph()
        # Title in bold
        run = p.add_run(f"{title} | {company} | {location}")
        run.font.size = Pt(11)
        run.font.bold = True
    
    def _add_dates(self, doc, dates):
        """Add date range in italic"""
        p = doc.add_paragraph()
        run = p.add_run(dates)
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(107, 114, 128)
    
    def _generate_professional_chronological(self, doc):
        """Generate Professional Chronological template content"""
        
        # Name
        self._add_name(doc)
        
        # Contact Info
        self._add_contact(doc, "your.email@example.com | +1-XXX-XXX-XXXX | City, State | LinkedIn: linkedin.com/in/yourprofile")
        
        # Professional Summary
        self._add_section_header(doc, "Professional Summary")
        self._add_placeholder(doc, 
            "Results-driven [Your Title] with X+ years of experience in [Industry/Field]. "
            "Proven track record of [Key Achievement]. Strong expertise in [Skill 1], [Skill 2], and [Skill 3]. "
            "Seeking to leverage technical skills and leadership experience to drive innovation at [Target Company].")
        
        # Professional Experience
        self._add_section_header(doc, "Professional Experience")
        
        # Job 1
        self._add_job_title(doc, "Job Title", "Company Name", "City, State")
        self._add_dates(doc, "Start Date – End Date (or Present)")
        self._add_bullet(doc, "Led/Managed/Developed [achievement] resulting in [X% improvement/$ savings/metric]")
        self._add_bullet(doc, "Implemented [solution] that increased [metric] by X% across Y projects")
        self._add_bullet(doc, "Collaborated with [team size] team members to deliver [project] serving X+ users")
        self._add_bullet(doc, "Reduced [process/cost] by X% through [action], saving $Y annually")
        
        doc.add_paragraph()  # Spacing
        
        # Job 2
        self._add_job_title(doc, "Previous Job Title", "Previous Company", "City, State")
        self._add_dates(doc, "Start Date – End Date")
        self._add_bullet(doc, "Developed [project/feature] using [technologies] for X+ users")
        self._add_bullet(doc, "Achieved [metric] by optimizing [process] through [method]")
        self._add_bullet(doc, "Mentored team of X engineers on [technology/practice]")
        
        # Technical Skills
        self._add_section_header(doc, "Technical Skills")
        p = doc.add_paragraph()
        run = p.add_run("Programming Languages: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = p.add_run("Python, Java, JavaScript, SQL, C++")
        run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        run = p.add_run("Frameworks & Tools: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = p.add_run("React, Node.js, Django, Docker, Kubernetes, Git")
        run.font.size = Pt(10)
        
        p = doc.add_paragraph()
        run = p.add_run("Cloud & Databases: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = p.add_run("AWS, Azure, PostgreSQL, MongoDB, Redis")
        run.font.size = Pt(10)
        
        # Education
        self._add_section_header(doc, "Education")
        self._add_job_title(doc, "Bachelor of Science in Computer Science", "University Name", "City, State")
        self._add_dates(doc, "Graduation Year | GPA: 3.X/4.0")
        self._add_bullet(doc, "Relevant Coursework: Data Structures, Algorithms, Machine Learning, Databases")
        
        # Certifications
        self._add_section_header(doc, "Certifications")
        self._add_bullet(doc, "AWS Certified Solutions Architect – Associate (Year)")
        self._add_bullet(doc, "Certified Kubernetes Administrator (CKA) (Year)")
    
    def _generate_modern_skills_focused(self, doc):
        """Generate Modern Skills-Focused template content"""
        
        # Name
        self._add_name(doc)
        
        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Software Engineer | Full Stack Developer")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Contact
        self._add_contact(doc, "your.email@example.com | +1-XXX-XXX-XXXX | github.com/yourusername | linkedin.com/in/yourprofile")
        
        # Technical Skills (prominent)
        self._add_section_header(doc, "Technical Skills")
        
        skills = [
            ("Languages:", "Python, JavaScript, TypeScript, Java, Go, SQL, HTML/CSS"),
            ("Frontend:", "React, Vue.js, Next.js, TailwindCSS, Redux, Material-UI"),
            ("Backend:", "Node.js, Express, FastAPI, Django, Spring Boot, REST APIs, GraphQL"),
            ("Databases:", "PostgreSQL, MongoDB, Redis, MySQL, Elasticsearch"),
            ("DevOps & Cloud:", "Docker, Kubernetes, AWS (EC2, S3, Lambda), Azure, CI/CD, GitHub Actions"),
            ("Tools:", "Git, VS Code, Postman, Jest, Pytest, Webpack, npm/yarn")
        ]
        
        for label, items in skills:
            p = doc.add_paragraph()
            run = p.add_run(label + " ")
            run.font.bold = True
            run.font.size = Pt(10)
            run = p.add_run(items)
            run.font.size = Pt(10)
        
        # Technical Projects
        self._add_section_header(doc, "Technical Projects")
        
        # Project 1
        p = doc.add_paragraph()
        run = p.add_run("Project Name")
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(" | ")
        run.font.size = Pt(10)
        run = p.add_run("Technologies: React, Node.js, PostgreSQL, AWS")
        run.italic = True
        run.font.size = Pt(10)
        
        self._add_dates(doc, "GitHub: github.com/yourusername/project | Live: yourproject.com")
        self._add_bullet(doc, "Built full-stack application serving X+ users with feature Y and Z")
        self._add_bullet(doc, "Implemented authentication system with JWT, achieving 99.9% uptime")
        self._add_bullet(doc, "Optimized database queries reducing load time by X%")
        
        doc.add_paragraph()  # Spacing
        
        # Project 2
        p = doc.add_paragraph()
        run = p.add_run("Another Project")
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(" | ")
        run.font.size = Pt(10)
        run = p.add_run("Technologies: Python, FastAPI, Docker, MongoDB")
        run.italic = True
        run.font.size = Pt(10)
        
        self._add_dates(doc, "GitHub: github.com/yourusername/project2")
        self._add_bullet(doc, "Developed RESTful API handling X requests/second with Y ms latency")
        self._add_bullet(doc, "Containerized application with Docker, deployed on AWS ECS")
        
        # Professional Experience
        self._add_section_header(doc, "Professional Experience")
        self._add_job_title(doc, "Software Engineer", "Company Name", "City, State")
        self._add_dates(doc, "Start Date – Present")
        self._add_bullet(doc, "Developed X features using [tech stack] impacting Y+ users")
        self._add_bullet(doc, "Improved system performance by X% through code optimization")
        self._add_bullet(doc, "Collaborated in Agile environment with team of X developers")
        
        # Education
        self._add_section_header(doc, "Education")
        self._add_job_title(doc, "Bachelor of Science in Computer Science", "University Name", "City, State")
        self._add_dates(doc, "Graduation Year | GPA: 3.X/4.0")
    
    def _generate_entry_level_student(self, doc):
        """Generate Entry-Level/Student template content"""
        
        # Name
        self._add_name(doc)
        
        # Contact
        self._add_contact(doc, "your.email@university.edu | +1-XXX-XXX-XXXX | linkedin.com/in/yourprofile | github.com/yourusername")
        
        # Education (first for students)
        self._add_section_header(doc, "Education")
        self._add_job_title(doc, "Bachelor of Science in Computer Science", "University Name", "City, State")
        self._add_dates(doc, "Expected Graduation: Month Year | GPA: 3.X/4.0")
        self._add_bullet(doc, "Relevant Coursework: Data Structures, Algorithms, Database Systems, Web Development, Software Engineering")
        self._add_bullet(doc, "Honors: Dean's List (Semester/Year), Academic Scholarship")
        
        # Technical Skills
        self._add_section_header(doc, "Technical Skills")
        
        skills = [
            ("Programming:", "Python, Java, JavaScript, C++, SQL, HTML/CSS"),
            ("Technologies:", "React, Node.js, Git, Docker, PostgreSQL, MongoDB"),
            ("Tools:", "VS Code, IntelliJ, Jupyter Notebook, Postman, GitHub")
        ]
        
        for label, items in skills:
            p = doc.add_paragraph()
            run = p.add_run(label + " ")
            run.font.bold = True
            run.font.size = Pt(10)
            run = p.add_run(items)
            run.font.size = Pt(10)
        
        # Projects
        self._add_section_header(doc, "Academic & Personal Projects")
        
        # Project 1
        p = doc.add_paragraph()
        run = p.add_run("Project Name")
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(" | ")
        run.font.size = Pt(10)
        run = p.add_run("Course: Software Engineering | Technologies: React, Python, SQL")
        run.italic = True
        run.font.size = Pt(10)
        
        self._add_dates(doc, "GitHub: github.com/yourusername/project")
        self._add_bullet(doc, "Developed web application for [purpose] as part of team project")
        self._add_bullet(doc, "Implemented [feature] using [technology], improving [metric]")
        self._add_bullet(doc, "Collaborated with team of X students using Agile methodology")
        
        doc.add_paragraph()  # Spacing
        
        # Project 2
        p = doc.add_paragraph()
        run = p.add_run("Another Project")
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(" | ")
        run.font.size = Pt(10)
        run = p.add_run("Personal Project | Technologies: Python, Flask, MongoDB")
        run.italic = True
        run.font.size = Pt(10)
        
        self._add_bullet(doc, "Built [application] to solve [problem] with X+ users/downloads")
        self._add_bullet(doc, "Integrated [API/service] to provide [functionality]")
        
        # Experience
        self._add_section_header(doc, "Experience")
        self._add_job_title(doc, "Software Engineering Intern", "Company Name", "City, State")
        self._add_dates(doc, "Month Year – Month Year")
        self._add_bullet(doc, "Contributed to [project/feature] using [technologies]")
        self._add_bullet(doc, "Wrote unit tests achieving X% code coverage")
        self._add_bullet(doc, "Participated in code reviews and daily standups")
        
        # Activities & Leadership
        self._add_section_header(doc, "Activities & Leadership")
        self._add_bullet(doc, "Computer Science Club – Member/Officer (Year-Year): Organized hackathons and tech talks")
        self._add_bullet(doc, "Hackathon Participant – Won [award] at [Hackathon Name] (Year)")
        self._add_bullet(doc, "Volunteer Tutor – Tutored X students in programming and mathematics")
    
    def get_template_info(self, template_type: str) -> Dict:
        """Get information about a template"""
        return self.TEMPLATES.get(template_type, {})
    
    def list_templates(self) -> List[Dict]:
        """List all available templates"""
        return [
            {
                'id': key,
                **value
            }
            for key, value in self.TEMPLATES.items()
        ]


# Test function
if __name__ == '__main__':
    generator = ResumeTemplateGenerator()
    
    # Generate all templates
    for template_id in ['professional_chronological', 'modern_skills_focused', 'entry_level_student']:
        output = f"/tmp/template_{template_id}.docx"
        success = generator.generate_template(template_id, output)
        print(f"{'✓' if success else '✗'} Generated: {output}")
